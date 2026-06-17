"""
Build a nicely formatted Word (.docx) document from REQUIREMENTS.md
Usage: python3 build_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Brand colours ────────────────────────────────────────────────────────────
BOOTS_BLUE   = RGBColor(0x05, 0x05, 0x4B)
BOOTS_LIGHT  = RGBColor(0x0C, 0x0C, 0x88)
HEADING_FG   = RGBColor(0x05, 0x05, 0x4B)
SUBHEAD_FG   = RGBColor(0x1E, 0x3A, 0x8A)
RED          = RGBColor(0xDC, 0x26, 0x26)
AMBER        = RGBColor(0xD9, 0x77, 0x06)
GREEN        = RGBColor(0x2E, 0x7D, 0x32)
CODE_BG      = RGBColor(0xF3, 0xF4, 0xF6)
CODE_FG      = RGBColor(0x1F, 0x2D, 0x3D)
TABLE_HEADER = RGBColor(0x05, 0x05, 0x4B)
TABLE_ALT    = RGBColor(0xF0, 0xF4, 0xFF)
MUTED        = RGBColor(0x6B, 0x72, 0x80)
DIVIDER      = RGBColor(0xE2, 0xE6, 0xEF)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

FONT_BODY = "Calibri"
FONT_CODE = "Courier New"
FONT_HEAD = "Calibri"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        if side in kwargs:
            border_el = OxmlElement(f'w:{side}')
            border_el.set(qn('w:val'), kwargs[side].get('val', 'single'))
            border_el.set(qn('w:sz'), str(kwargs[side].get('sz', 4)))
            border_el.set(qn('w:color'), kwargs[side].get('color', 'auto'))
            tcBorders.append(border_el)
    tcPr.append(tcBorders)


def hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_page_number(doc: Document):
    """Add page numbers to the footer."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    # Add "Page X of Y"
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run2 = p.add_run(' of ')
    run2.font.size = Pt(9)
    run2.font.color.rgb = MUTED

    run3 = p.add_run()
    run3.font.size = Pt(9)
    run3.font.color.rgb = MUTED
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar3)
    run3._r.append(instrText2)
    run3._r.append(fldChar4)


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    add_page_break_run(run)


def add_page_break_run(run):
    """Append a page-break XML element directly to a run."""
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def set_para_spacing(para, before: int = 0, after: int = 0, line: float = 1.0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def rgb_hex(rgb: RGBColor) -> str:
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'


def add_colored_bar(doc: Document, color: RGBColor = BOOTS_BLUE, height_pts: int = 3):
    """Add a thin horizontal rule as a 1-row, 1-col table with colored background."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    # Set row height
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pts * 20)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)
    # Remove borders
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is not None:
        tblBorders = OxmlElement('w:tblBorders')
        for side in ['top','left','bottom','right','insideH','insideV']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tblBorders.append(el)
        tblPr.append(tblBorders)
    doc.add_paragraph()


# ─── Inline markdown parser ────────────────────────────────────────────────────
def apply_inline(para, text: str, base_size: Pt = Pt(11), base_color: RGBColor = None,
                 mono: bool = False, bold_base: bool = False):
    """
    Parse inline markdown: **bold**, `code`, ~~strike~~, [link](url) → underline
    and add runs to para.
    """
    # Pattern order matters
    pattern = re.compile(r'(\*\*(.+?)\*\*|`([^`]+)`|\[([^\]]+)\]\([^\)]+\))')
    pos = 0
    for m in pattern.finditer(text):
        # Text before match
        before = text[pos:m.start()]
        if before:
            run = para.add_run(before)
            run.font.size = base_size
            run.font.name = FONT_CODE if mono else FONT_BODY
            if base_color:
                run.font.color.rgb = base_color
            if bold_base:
                run.font.bold = True

        full = m.group(0)
        if full.startswith('**'):
            # Bold
            run = para.add_run(m.group(2))
            run.font.bold = True
            run.font.size = base_size
            run.font.name = FONT_CODE if mono else FONT_BODY
            if base_color:
                run.font.color.rgb = base_color
        elif full.startswith('`'):
            # Inline code
            run = para.add_run(m.group(3))
            run.font.name = FONT_CODE
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        elif full.startswith('['):
            # Link text only (no hyperlink in docx for simplicity)
            run = para.add_run(m.group(4))
            run.font.underline = True
            run.font.color.rgb = BOOTS_LIGHT
            run.font.size = base_size

        pos = m.end()

    # Remaining text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.size = base_size
        run.font.name = FONT_CODE if mono else FONT_BODY
        if base_color:
            run.font.color.rgb = base_color
        if bold_base:
            run.font.bold = True


# ─── Cover page ───────────────────────────────────────────────────────────────
def build_cover(doc: Document):
    # Full-width blue bar at top (via a table trick)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, BOOTS_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Set row height to 220pt
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(220 * 20))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)
    # Remove all borders
    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is not None:
        tblBorders = OxmlElement('w:tblBorders')
        for side in ['top','left','bottom','right','insideH','insideV']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tblBorders.append(el)
        tblPr.append(tblBorders)

    # Service labels inside the blue box
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run('BLEAF')
    run.font.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = WHITE
    run.font.name = FONT_HEAD

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Clinical Workforce Management Tool')
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0xB8, 0xC4, 0xE8)
    run2.font.name = FONT_HEAD

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Boots Digital Health')
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x80, 0x90, 0xC0)
    run3.font.name = FONT_HEAD

    # Below the blue box
    doc.add_paragraph()
    doc.add_paragraph()

    meta_fields = [
        ('Document type', 'Product Requirements Document'),
        ('Version', '1.0  —  Prototype → Production'),
        ('Date', 'June 2026'),
        ('Technology', 'TypeScript  ·  Node.js  ·  Python'),
        ('Audience', 'Development team, QA, Product'),
    ]
    tbl2 = doc.add_table(rows=len(meta_fields), cols=2)
    tbl2.style = 'Table Grid'
    for i, (label, value) in enumerate(meta_fields):
        row = tbl2.rows[i]
        # Label cell
        lc = row.cells[0]
        set_cell_bg(lc, RGBColor(0xF0, 0xF4, 0xFF))
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run(label)
        lr.font.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = BOOTS_BLUE
        lr.font.name = FONT_BODY
        lc.width = Inches(2)

        # Value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after = Pt(4)
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
        vr.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        vr.font.name = FONT_BODY

    # Remove table borders to keep it clean
    tbl2_xml = tbl2._tbl
    tblPr2 = tbl2_xml.find(qn('w:tblPr'))
    if tblPr2 is None:
        tblPr2 = OxmlElement('w:tblPr')
        tbl2_xml.insert(0, tblPr2)
    tblBorders2 = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '2')
        el.set(qn('w:color'), rgb_hex(DIVIDER))
        tblBorders2.append(el)
    tblPr2.append(tblBorders2)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run('Generated from working prototype  ·  All code snippets are TypeScript, Python, or Node.js')
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    r.font.italic = True

    # Page break after cover
    pb = doc.add_paragraph()
    add_page_break_run(pb.add_run())


# ─── Main document builder ─────────────────────────────────────────────────────
def build_doc(md_path: str, output_path: str):
    doc = Document()

    # ── Page setup
    section = doc.sections[0]
    section.page_width  = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # ── Base Normal style
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # ── Heading styles
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = FONT_HEAD
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.color.rgb = WHITE
    h1_style.paragraph_format.space_before = Pt(0)
    h1_style.paragraph_format.space_after  = Pt(4)
    h1_style.paragraph_format.keep_with_next = True

    h2_style = doc.styles['Heading 2']
    h2_style.font.name = FONT_HEAD
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = BOOTS_BLUE
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after  = Pt(4)
    h2_style.paragraph_format.keep_with_next = True

    h3_style = doc.styles['Heading 3']
    h3_style.font.name = FONT_HEAD
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = SUBHEAD_FG
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after  = Pt(3)
    h3_style.paragraph_format.keep_with_next = True

    add_page_number(doc)

    # ── Cover page
    build_cover(doc)

    # ── Read and parse markdown
    md_text = Path(md_path).read_text(encoding='utf-8')
    lines = md_text.split('\n')

    i = 0
    current_epic_color = BOOTS_BLUE

    # Colours to cycle through for epic banners
    epic_colors = [
        RGBColor(0xC2, 0x18, 0x5B),  # Women's pink
        RGBColor(0x15, 0x65, 0xC0),  # Men's blue
        RGBColor(0x2E, 0x7D, 0x32),  # Green
        RGBColor(0x6A, 0x1B, 0x9A),  # Purple
        RGBColor(0xE6, 0x51, 0x00),  # Orange
        RGBColor(0x00, 0x83, 0x8F),  # Teal
        RGBColor(0xD9, 0x77, 0x06),  # Amber
        RGBColor(0x05, 0x05, 0x4B),  # Boots blue
    ]
    epic_idx = 0

    while i < len(lines):
        line = lines[i]

        # ── Skip raw HTML / horizontal rules / table of contents entries
        if line.strip().startswith('<') or line.strip() == '---':
            if line.strip() == '---':
                doc.add_paragraph()  # small spacer
            i += 1
            continue

        # ── H1 (document title — treated as section header)
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            # Only show if it's not the very first H1 (which is on the cover)
            if 'BLeaf Clinical' not in title_text:
                color = epic_colors[epic_idx % len(epic_colors)]
                epic_idx += 1
                current_epic_color = color

                # Blue banner paragraph
                tbl = doc.add_table(rows=1, cols=1)
                cell = tbl.cell(0, 0)
                set_cell_bg(cell, color)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(title_text)
                run.font.bold  = True
                run.font.size  = Pt(18)
                run.font.color.rgb = WHITE
                run.font.name  = FONT_HEAD
                # Remove borders
                tbl_xml = tbl._tbl
                tblPr = tbl_xml.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl_xml.insert(0, tblPr)
                tblBorders = OxmlElement('w:tblBorders')
                for side in ['top','left','bottom','right','insideH','insideV']:
                    el = OxmlElement(f'w:{side}')
                    el.set(qn('w:val'), 'none')
                    tblBorders.append(el)
                tblPr.append(tblBorders)
                doc.add_paragraph()
            i += 1
            continue

        # ── H2
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            p = doc.add_paragraph(style='Heading 2')
            apply_inline(p, text, base_size=Pt(14), base_color=BOOTS_BLUE)
            set_para_spacing(p, before=20, after=4)
            i += 1
            continue

        # ── H3
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph(style='Heading 3')
            apply_inline(p, text, base_size=Pt(12), base_color=SUBHEAD_FG)
            set_para_spacing(p, before=14, after=3)
            i += 1
            continue

        # ── H4
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            set_para_spacing(p, before=10, after=2)
            i += 1
            continue

        # ── Code block
        if line.strip().startswith('```'):
            lang = line.strip()[3:]
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            # Render as table with coloured header + code body
            code_text = '\n'.join(code_lines)

            # Header row with language label
            tbl = doc.add_table(rows=2, cols=1)
            tbl_xml = tbl._tbl
            tblPr = tbl_xml.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl_xml.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '2')
                el.set(qn('w:color'), rgb_hex(RGBColor(0xD1, 0xD5, 0xDB)))
                tblBorders.append(el)
            tblPr.append(tblBorders)

            # Header cell (language tag)
            hdr_cell = tbl.cell(0, 0)
            set_cell_bg(hdr_cell, RGBColor(0x1F, 0x29, 0x37))
            hdr_p = hdr_cell.paragraphs[0]
            hdr_p.paragraph_format.space_before = Pt(2)
            hdr_p.paragraph_format.space_after  = Pt(2)
            hdr_run = hdr_p.add_run(lang.upper() if lang else 'CODE')
            hdr_run.font.name  = FONT_CODE
            hdr_run.font.size  = Pt(8)
            hdr_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            hdr_run.font.bold  = True

            # Code cell
            code_cell = tbl.cell(1, 0)
            set_cell_bg(code_cell, RGBColor(0xF8, 0xF9, 0xFA))
            # Clear existing paragraph then add lines
            for para in code_cell.paragraphs:
                for run in para.runs:
                    run.text = ''
            # Add the code text, splitting by newline
            first = True
            for cl in code_lines:
                if first:
                    cp = code_cell.paragraphs[0]
                    first = False
                else:
                    cp = code_cell.add_paragraph()
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after  = Pt(0)
                # Syntax-colour keywords for TypeScript / Python / SQL
                cr = cp.add_run(cl)
                cr.font.name  = FONT_CODE
                cr.font.size  = Pt(8.5)
                cr.font.color.rgb = CODE_FG

            # Add cell padding
            for cell in [hdr_cell, code_cell]:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for side in ['top','left','bottom','right']:
                    m = OxmlElement(f'w:{side}')
                    m.set(qn('w:w'), '80' if side in ['left','right'] else '40')
                    m.set(qn('w:type'), 'dxa')
                    tcMar.append(m)
                tcPr.append(tcMar)

            # Width to full page
            tblPr_width = OxmlElement('w:tblW')
            tblPr_width.set(qn('w:w'), '9360')
            tblPr_width.set(qn('w:type'), 'dxa')
            tblPr.append(tblPr_width)

            doc.add_paragraph()
            continue

        # ── Markdown table
        if line.strip().startswith('|') and '|' in line[1:]:
            # Collect all table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Filter out separator rows (|---|---|)
            data_rows = [r for r in table_lines if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', r)]
            if not data_rows:
                continue

            # Parse cells
            parsed = []
            for row_line in data_rows:
                cells = [c.strip() for c in row_line.strip().strip('|').split('|')]
                parsed.append(cells)

            if not parsed:
                continue

            num_cols = max(len(r) for r in parsed)
            # Pad rows to same width
            for r in parsed:
                while len(r) < num_cols:
                    r.append('')

            tbl = doc.add_table(rows=len(parsed), cols=num_cols)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            for ri, row_data in enumerate(parsed):
                is_header = (ri == 0)
                row = tbl.rows[ri]
                for ci, cell_text in enumerate(row_data):
                    cell = row.cells[ci]
                    if is_header:
                        set_cell_bg(cell, BOOTS_BLUE)
                    elif ri % 2 == 0:
                        set_cell_bg(cell, RGBColor(0xF8, 0xFA, 0xFF))

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    if is_header:
                        apply_inline(p, cell_text, base_size=Pt(9.5),
                                     base_color=WHITE, bold_base=True)
                    else:
                        apply_inline(p, cell_text, base_size=Pt(9.5))

                    # Cell padding
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcMar = OxmlElement('w:tcMar')
                    for side in ['top','left','bottom','right']:
                        m = OxmlElement(f'w:{side}')
                        m.set(qn('w:w'), '80')
                        m.set(qn('w:type'), 'dxa')
                        tcMar.append(m)
                    tcPr.append(tcMar)

            # Style table borders
            tbl_xml = tbl._tbl
            tblPr = tbl_xml.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl_xml.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '2')
                el.set(qn('w:color'), rgb_hex(RGBColor(0xD1, 0xD5, 0xDB)))
                tblBorders.append(el)
            tblPr.append(tblBorders)

            doc.add_paragraph()
            continue

        # ── Bullet list items
        if re.match(r'^[-*] ', line) or re.match(r'^\s+[-*] ', line):
            indent = len(line) - len(line.lstrip())
            raw = re.sub(r'^[\s\-\*]+', '', line).strip()

            # Checkbox item
            is_check = raw.startswith('[ ]') or raw.startswith('[x]') or raw.startswith('[X]')
            if is_check:
                checked = raw[1].lower() == 'x'
                raw = raw[3:].strip()
                p = doc.add_paragraph(style='List Bullet')
                prefix_run = p.add_run('☑ ' if checked else '☐ ')
                prefix_run.font.color.rgb = GREEN if checked else MUTED
                prefix_run.font.size = Pt(10.5)
                apply_inline(p, raw, base_size=Pt(10.5))
            else:
                p = doc.add_paragraph(style='List Bullet')
                apply_inline(p, raw, base_size=Pt(10.5))

            p.paragraph_format.left_indent  = Inches(0.3 + indent * 0.02)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # ── Numbered list
        if re.match(r'^\d+\. ', line) or re.match(r'^\s+\d+\. ', line):
            raw = re.sub(r'^[\s\d\.]+', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            apply_inline(p, raw, base_size=Pt(10.5))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            i += 1
            continue

        # ── Blockquote (> text)
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            # Left border via table
            tbl = doc.add_table(rows=1, cols=2)
            tbl_xml = tbl._tbl
            tblPr = tbl_xml.find(qn('w:tblPr'))
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl_xml.insert(0, tblPr)
            tblBorders = OxmlElement('w:tblBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tblBorders.append(el)
            tblPr.append(tblBorders)

            # Colour bar col
            bar_cell = tbl.cell(0, 0)
            set_cell_bg(bar_cell, RGBColor(0x93, 0xC5, 0xFD))
            bar_cell.width = Inches(0.08)
            # Set fixed width
            tc = bar_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '110')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

            text_cell = tbl.cell(0, 1)
            set_cell_bg(text_cell, RGBColor(0xEF, 0xF6, 0xFF))
            tp = text_cell.paragraphs[0]
            tp.paragraph_format.space_before = Pt(5)
            tp.paragraph_format.space_after  = Pt(5)
            apply_inline(tp, text, base_size=Pt(10.5),
                         base_color=RGBColor(0x1E, 0x40, 0xAF))
            doc.add_paragraph()
            # Remove the unused paragraph we created earlier
            # (it got added before we decided to use a table)
            elem = p._element
            elem.getparent().remove(elem)
            i += 1
            continue

        # ── User story heading (US-XXX-NN pattern)
        if re.match(r'\*\*US-[A-Z]+-\d+\*\*', line.strip()):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            set_para_spacing(p, before=12, after=4)
            r = p.add_run(text)
            r.font.bold  = True
            r.font.size  = Pt(11)
            r.font.color.rgb = BOOTS_BLUE
            r.font.name  = FONT_BODY
            i += 1
            continue

        # ── Empty line
        if not line.strip():
            # Add a small gap but not a full paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            i += 1
            continue

        # ── Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)
        apply_inline(p, line.strip(), base_size=Pt(11))
        i += 1

    # ── Save
    doc.save(output_path)
    print(f'✅  Saved: {output_path}')


if __name__ == '__main__':
    build_doc(
        md_path='/home/user/bleaf-demo/REQUIREMENTS.md',
        output_path='/home/user/bleaf-demo/BLeaf_Requirements.docx',
    )
